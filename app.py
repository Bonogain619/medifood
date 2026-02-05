import streamlit as st
from google import genai
from google.genai import types
import os, re, io
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# 1. 시스템 상수 및 API 설정
MODEL_NAME = "gemini-3-flash-preview"

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise RuntimeError("GEMINI_API_KEY 환경변수가 설정되지 않았습니다.")
MODEL_NAME = "gemini-3-flash-preview"

st.set_page_config(page_title="메디푸드 분석 시스템", layout="wide", page_icon="🔬")

DISEASE_LIST = [
    "고혈압", "당뇨(1형/2형)", "고지혈증", "신장질환(CKD)", "투석 중",
    "간경화/지방간", "위염/위궤양", "역류성 식도염", "크론병/궤양성대장염",
    "갑상선 질환", "통풍", "골다공증", "심부전", "암 관리", "빈혈", "비만"
]

# 2. Word(DOCX) 생성 엔진 (v3.1 고정 레이아웃 버전)
def set_korean_font(run):
    run.font.name = '맑은 고딕'
    r = run._element
    rFonts = r.get_or_add_rPr().get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')

def create_docx_report(content):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '맑은 고딕')

    title = doc.add_heading('메디푸드 분석 리포트', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    lines = content.split('\n')
    table_data, is_table = [], False
    
    for line in lines:
        if re.search(r'\|.*\|', line):
            is_table = True
            row = [re.sub(r'\*\*', '', c.strip()) for c in line.split('|') if c.strip() and not set(c.strip()).issubset({'-', ':', '|'})]
            if row: table_data.append(row)
        else:
            if is_table and table_data:
                # [핵심] 식단표 너비 고정 로직 적용
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                table.autofit = False # 자동 조정 비활성화
                
                # 컬럼 너비 인치 단위 강제 고정 (A4 가로폭 최적화)
                # 구분(1.0), 아침/점심/저녁/간식(각 1.5) 등 컬럼 수에 따라 분배
                num_cols = len(table_data[0])
                total_w = 7.0 # 가용 너비 7인치
                col_w = total_w / num_cols

                for i, row_data in enumerate(table_data):
                    for j, cell_text in enumerate(row_data):
                        cell = table.cell(i, j)
                        cell.width = Inches(col_w)
                        
                        paragraph = cell.paragraphs[0]
                        run = paragraph.add_run(cell_text)
                        set_korean_font(run)
                        
                        if i == 0 or j == 0:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            
                doc.add_paragraph()
                table_data, is_table = [], False
            
            clean_line = re.sub(r'[\*\#]', '', line)
            if clean_line.strip():
                p = doc.add_paragraph()
                run = p.add_run(clean_line.strip())
                set_korean_font(run)
    return doc

# 3. 세션 및 프로세스 관리
if "session_id" not in st.session_state:
    st.session_state.session_id = 0
if "analysis_result" not in st.session_state:
    st.session_state.analysis_result = ""

def reset_system():
    st.session_state.analysis_result = ""
    st.session_state.session_id += 1
    st.rerun()

def shutdown_app():
    st.warning("시스템을 종료합니다. 서버 프로세스가 중단됩니다.")
    os._exit(0)

# --- UI 레이아웃 (v2.9 원본 폼 구조 복구) ---
st.title("🔬 메디푸드 분석 시스템")

with st.sidebar:
    st.header("⚙️ 시스템 관리")
    if st.button("🔄 새 상담 시작 (데이터 리셋)"):
        reset_system()
    
    st.success("API Key 자동 인증 활성화")
    
    st.divider()
    st.header("📋 데이터 입력")
    s_id = st.session_state.session_id
    age = st.number_input("나이", min_value=1, value=None, placeholder="나이 입력", key=f"age_{s_id}")
    gender = st.radio("성별", ["남성", "여성"], key=f"gen_{s_id}")
    disease = st.multiselect("기저질환", DISEASE_LIST, key=f"dis_{s_id}")
    # [복구] 사이드바 복용 약물 입력란
    medication = st.text_input("복용 중인 약물", key=f"med_{s_id}")

    st.sidebar.markdown("<br><br>" * 5, unsafe_allow_html=True)
    st.divider()
    if st.button("🔴 시스템 종료"):
        shutdown_app()

# [복구] 본문 증상 입력 및 분석 실행 버튼
symptom = st.text_area("현재 증상 및 상세 특징", height=150, key=f"sym_{s_id}")

if st.button("🚀 정밀 분석 및 식단표 생성"):
    if age is None or not symptom:
        st.warning("분석을 위해 나이와 증상을 입력해 주십시오.")
    else:
        try:
            client = genai.Client(api_key=FIXED_API_KEY)
            prompt = f"""
            [Role] 당신은 '메디푸드 분석 시스템'입니다.
            [Instruction] 약초 제외. 30일 식단표는 반드시 각 주차별로 표로 작성.
            [User Data] 나이:{age}, 성별:{gender}, 질환:{disease}, 약물:{medication}, 증상:{symptom}
            """
            with st.spinner("AI 엔진이 리포트 서식을 최적화 중입니다..."):
                response = client.models.generate_content(model=MODEL_NAME, contents=prompt)
                st.session_state.analysis_result = response.text
        except Exception as e:
            st.error(f"실행 오류: {e}")

# 5. 결과 출력 및 Word 다운로드
if st.session_state.analysis_result:
    st.divider()
    col_l, col_r = st.columns([8, 2])
    with col_l: st.subheader("📋 메디푸드 정밀 분석 결과")
    with col_r:
        doc_obj = create_docx_report(st.session_state.analysis_result)
        doc_stream = io.BytesIO()
        doc_obj.save(doc_stream)
        doc_stream.seek(0)
        st.download_button(
            label="📥 Word 리포트 저장",
            data=doc_stream,
            file_name=f"medifood_report_{datetime.now().strftime('%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    st.markdown(st.session_state.analysis_result)
